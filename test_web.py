import io
import re
import random
import colorsys
import html as html_module
import json
import os
from collections import Counter
from typing import List, Dict, Set
from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import matplotlib.pyplot as plt
import matplotlib
import math
matplotlib.use('Agg')
from datetime import datetime
from wordcloud import WordCloud

app = FastAPI(title="N-gram Chinese Text Analyzer", version="2.0")

# 정적 파일 및 템플릿 설정
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# 다국어 로케일 (하드코딩)
LOCALES = {
    "ko": {
        "fileUpload": "파일 업로드",
        "minNgram": "N-gram 최소 값(2부터)",
        "maxNgram": "N-gram 최대 값(예: 4)",
        "analyze": "분석 시작",
        "wordcloud": "워드클라우드",
        "download": "다운로드",
        "applyHighlight": "하이라이트 적용",
        "mode": "모드",
        "all": "전체 분석",
        "common": "공통 부분",
        "selectCommon": "공통 부분 선택",
        "includeAll": "모든 교집합 포함",
        "sort": "정렬 기준",
        "freqAsc": "빈도수 오름차순",
        "freqDesc": "빈도수 내림차순",
        "dataCount": "데이터 항목 수: {count}개",
        "needTwo": "최소 2개 이상의 파일을 선택해야 합니다.",
        "noData": "먼저 분석을 수행하세요.",
        "analysisControls": "분석 제어",
        "results": "분석 결과",
        "processing": "처리 중...",
        "pleaseWait": "잠시만 기다려 주세요.",
        "selectFileForHighlight": "하이라이트 적용할 파일 선택",
        "cancel": "취소",
        "apply": "적용",
        "downloadOptions": "다운로드 옵션",
        "downloadTxt": "텍스트 파일 (.txt)",
        "downloadHtml": "HTML 파일 (.html)",
        "error": "오류",
        "success": "성공",
        "fileReadError": "파일 읽기 오류",
        "analysisError": "분석 중 오류가 발생했습니다.",
        "wordcloudError": "워드클라우드 생성 중 오류가 발생했습니다.",
        "highlightError": "하이라이트 적용 중 오류가 발생했습니다.",
        "downloadError": "다운로드 중 오류가 발생했습니다.",
        "noValidData": "유효한 데이터가 없습니다.",
        "selectFiles": "파일을 선택하세요",
        "analyzing": "분석 중...",
        "generatingWordcloud": "워드클라우드 생성 중...",
        "applyingHighlight": "하이라이트 적용 중...",
        "downloading": "다운로드 중...",
        "ngramRangeError": "N-gram 범위가 올바르지 않습니다.",
        "analysisComplete": "분석이 완료되었습니다.",
        "filterError": "필터링 중 오류가 발생했습니다.",
        "analyzeFirst": "먼저 파일을 분석해주세요.",
        "selectFile": "파일을 선택해주세요.",
        "selectedFilesCount": "선택된 파일:"
    },
    "en": {
        "fileUpload": "Upload Files",
        "minNgram": "Min N-gram (from 2)",
        "maxNgram": "Max N-gram (e.g., 4)",
        "analyze": "Start Analysis",
        "wordcloud": "Word Cloud",
        "download": "Download",
        "applyHighlight": "Apply Highlight",
        "mode": "Mode",
        "all": "All Analysis",
        "common": "Common Parts",
        "selectCommon": "Select Common Parts",
        "includeAll": "Include All Intersections",
        "sort": "Sort By",
        "freqAsc": "Frequency Ascending",
        "freqDesc": "Frequency Descending",
        "dataCount": "Data Items: {count}",
        "needTwo": "You must select at least two files.",
        "noData": "Please run the analysis first.",
        "analysisControls": "Analysis Controls",
        "results": "Analysis Results",
        "processing": "Processing...",
        "pleaseWait": "Please wait...",
        "selectFileForHighlight": "Select File for Highlight",
        "cancel": "Cancel",
        "apply": "Apply",
        "downloadOptions": "Download Options",
        "downloadTxt": "Text File (.txt)",
        "downloadHtml": "HTML File (.html)",
        "error": "Error",
        "success": "Success",
        "fileReadError": "File read error",
        "analysisError": "Analysis error occurred.",
        "wordcloudError": "Wordcloud generation error occurred.",
        "highlightError": "Highlight application error occurred.",
        "downloadError": "Download error occurred.",
        "noValidData": "No valid data.",
        "selectFiles": "Select files",
        "analyzing": "Analyzing...",
        "generatingWordcloud": "Generating wordcloud...",
        "applyingHighlight": "Applying highlight...",
        "downloading": "Downloading...",
        "ngramRangeError": "N-gram range is invalid.",
        "analysisComplete": "Analysis completed.",
        "filterError": "Filtering error occurred.",
        "analyzeFirst": "Please analyze files first.",
        "selectFile": "Please select a file.",
        "selectedFilesCount": "Selected files:"
    },
    "zh": {
        "fileUpload": "上传文件",
        "minNgram": "最小N-gram（从2开始）",
        "maxNgram": "最大N-gram（例如：4）",
        "analyze": "开始分析",
        "wordcloud": "词云",
        "download": "下载",
        "applyHighlight": "应用高亮",
        "mode": "模式",
        "all": "全部分析",
        "common": "公共部分",
        "selectCommon": "选择公共部分",
        "includeAll": "包含所有交集",
        "sort": "排序方式",
        "freqAsc": "频率升序",
        "freqDesc": "频率降序",
        "dataCount": "数据项数：{count}",
        "needTwo": "必须选择至少两个文件。",
        "noData": "请先运行分析。",
        "analysisControls": "分析控制",
        "results": "分析结果",
        "processing": "处理中...",
        "pleaseWait": "请稍候...",
        "selectFileForHighlight": "选择要高亮的文件",
        "cancel": "取消",
        "apply": "应用",
        "downloadOptions": "下载选项",
        "downloadTxt": "文本文件 (.txt)",
        "downloadHtml": "HTML文件 (.html)",
        "error": "错误",
        "success": "成功",
        "fileReadError": "文件读取错误",
        "analysisError": "分析过程中发生错误。",
        "wordcloudError": "词云生成过程中发生错误。",
        "highlightError": "高亮应用过程中发生错误。",
        "downloadError": "下载过程中发生错误。",
        "noValidData": "没有有效数据。",
        "selectFiles": "选择文件",
        "analyzing": "分析中...",
        "generatingWordcloud": "生成词云中...",
        "applyingHighlight": "应用高亮中...",
        "downloading": "下载中...",
        "ngramRangeError": "N-gram范围无效。",
        "analysisComplete": "分析完成。",
        "filterError": "过滤过程中发生错误。",
        "analyzeFirst": "请先分析文件。",
        "selectFile": "请选择文件。",
        "selectedFilesCount": "已选择文件:"
    },
    "ja": {
        "fileUpload": "ファイルをアップロード",
        "minNgram": "最小 N-gram（2から）",
        "maxNgram": "最大 N-gram（例：4）",
        "analyze": "分析開始",
        "wordcloud": "ワードクラウド",
        "download": "ダウンロード",
        "applyHighlight": "ハイライト適用",
        "mode": "モード",
        "all": "全体分析",
        "common": "共通部分",
        "selectCommon": "共通部分を選択",
        "includeAll": "すべての共通項を含む",
        "sort": "ソート基準",
        "freqAsc": "頻度昇順",
        "freqDesc": "頻度降順",
        "dataCount": "データ項目数：{count}",
        "needTwo": "ファイルを最低2つ選択してください。",
        "noData": "まず分析を実行してください。",
        "analysisControls": "分析制御",
        "results": "分析結果",
        "processing": "処理中...",
        "pleaseWait": "お待ちください...",
        "selectFileForHighlight": "ハイライト適用するファイルを選択",
        "cancel": "キャンセル",
        "apply": "適用",
        "downloadOptions": "ダウンロードオプション",
        "downloadTxt": "テキストファイル (.txt)",
        "downloadHtml": "HTMLファイル (.html)",
        "error": "エラー",
        "success": "成功",
        "fileReadError": "ファイル読み込みエラー",
        "analysisError": "分析中にエラーが発生しました。",
        "wordcloudError": "ワードクラウド生成中にエラーが発生しました。",
        "highlightError": "ハイライト適用中にエラーが発生しました。",
        "downloadError": "ダウンロード中にエラーが発生しました。",
        "noValidData": "有効なデータがありません。",
        "selectFiles": "ファイルを選択",
        "analyzing": "分析中...",
        "generatingWordcloud": "ワードクラウド生成中...",
        "applyingHighlight": "ハイライト適用中...",
        "downloading": "ダウンロード中...",
        "ngramRangeError": "N-gram範囲が無効です。",
        "analysisComplete": "分析が完了しました。",
        "filterError": "フィルタリング中にエラーが発生しました。",
        "analyzeFirst": "まずファイルを分析してください。",
        "selectFile": "ファイルを選択してください。",
        "selectedFilesCount": "選択されたファイル:"
    }
}

# 한자 패턴 (N-gram.py와 동일)
HANJA_PATTERN = re.compile(r'[\u4e00-\u9fff\u3400-\u4DBF\uF900-\uFAFF\U00020000-\U0002A6DF]')

class NgramAnalyzer:
    """N-gram 분석 클래스 (원본 N-gram.py의 핵심 로직)"""
    
    def __init__(self):
        self.files_data = {}           # 파일명 -> 한자만 남긴 텍스트
        self.index_map = {}            # 파일명 -> 원본 텍스트에서 한자 위치(인덱스) 리스트
        self.original_text_data = {}   # 파일명 -> 원본 전체 텍스트
        self.ngram_results = {}        # 파일명 -> n-gram 카운터
        self.all_ngrams = set()        # 전체 n-gram 집합
    
    def process_files(self, files: List[UploadFile], min_n: int, max_n: int):
        """파일들을 처리하여 N-gram 분석 수행"""
        self.files_data.clear()
        self.index_map.clear()
        self.original_text_data.clear()
        self.ngram_results.clear()
        self.all_ngrams.clear()
        
        for file in files:
            try:
                content = file.file.read().decode("utf-8", errors="ignore")
                file.file.seek(0)  # 파일 포인터 리셋
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"파일 {file.filename} 읽기 오류: {str(e)}")
            
            file_name = os.path.splitext(file.filename)[0]
            self.original_text_data[file_name] = content
            
            # 한자만 추출
            processed_text = ""
            index_mapping = []
            for i, ch in enumerate(content):
                if HANJA_PATTERN.match(ch):
                    processed_text += ch
                    index_mapping.append(i)
            
            self.files_data[file_name] = processed_text
            self.index_map[file_name] = index_mapping
            
            # N-gram 카운팅
            counter = Counter()
            text_len = len(processed_text)
            for n in range(min_n, max_n + 1):
                for i in range(text_len - n + 1):
                    ngram = processed_text[i:i+n]
                    counter[ngram] += 1
            
            self.ngram_results[file_name] = counter
            self.all_ngrams.update(counter.keys())
    
    def get_analysis_results(self, sort_option: str = "빈도수 오름차순"):
        """분석 결과 반환"""
        file_keys = list(self.ngram_results.keys())
        sorted_ngrams = self.sort_ngrams(self.all_ngrams, file_keys, sort_option)
        
        lines = []
        for ngram in sorted_ngrams:
            parts = []
            for fname, counter in self.ngram_results.items():
                count = counter.get(ngram, 0)
                parts.append(f"{fname}: {count}")
            lines.append(f"{ngram}: " + ", ".join(parts))
        
        return lines
    
    def get_common_ngrams(self, selected_files: List[str], include_all_common: bool = False):
        """공통 N-gram 필터링"""
        if not selected_files:
            return []
        
        filtered_ngrams = []
        for ngram in self.all_ngrams:
            if include_all_common:
                # 모든 선택된 파일에 나타나는 N-gram
                if all(self.ngram_results[fname].get(ngram, 0) > 0 for fname in selected_files):
                    # 모든 파일의 빈도수 포함
                    parts = []
                    for fname in self.ngram_results.keys():
                        count = self.ngram_results[fname].get(ngram, 0)
                        parts.append(f"{fname}: {count}")
                    filtered_ngrams.append(f"{ngram}: " + ", ".join(parts))
            else:
                # 정확히 선택된 파일들에만 나타나는 N-gram
                selected_set = set(selected_files)
                appearance_set = {fname for fname, counter in self.ngram_results.items() if counter.get(ngram, 0) > 0}
                if appearance_set == selected_set and appearance_set:
                    # 모든 파일의 빈도수 포함
                    parts = []
                    for fname in self.ngram_results.keys():
                        count = self.ngram_results[fname].get(ngram, 0)
                        parts.append(f"{fname}: {count}")
                    filtered_ngrams.append(f"{ngram}: " + ", ".join(parts))
        
        return filtered_ngrams
    
    def sort_ngrams(self, ngrams, file_set, sort_option: str):
        """N-gram 정렬"""
        def freq_key(ngram):
            return sum(self.ngram_results[fname].get(ngram, 0) for fname in file_set)
        
        reverse = (sort_option == "빈도수 내림차순")
        return sorted(ngrams, key=freq_key, reverse=reverse)
    
    def apply_to_original_text(self, file_name: str, selected_ngrams: Set[str], lang: str = "ko"):
        """원본 텍스트에 선택된 N-gram 하이라이트 적용"""
        if file_name not in self.original_text_data:
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        orig_text = self.original_text_data[file_name]
        processed_text = self.files_data[file_name]
        index_mapping = self.index_map[file_name]
        
        # 하이라이트할 위치 찾기
        red_positions = set()
        for ngram in selected_ngrams:
            start_idx = 0
            while True:
                pos = processed_text.find(ngram, start_idx)
                if pos == -1:
                    break
                for offset in range(len(ngram)):
                    if pos + offset < len(index_mapping):
                        red_positions.add(index_mapping[pos + offset])
                start_idx = pos + 1
        
        # 언어별 텍스트
        lang_texts = {
            "ko": {
                "title": "분석 결과",
                "total_chars": "전체 한자수",
                "highlighted_chars": "해당 글자수",
                "ratio": "비율",
                "download": "다운로드",
                "back": "뒤로 가기"
            },
            "en": {
                "title": "Analysis Results",
                "total_chars": "Total Hanja Count",
                "highlighted_chars": "Highlighted Characters",
                "ratio": "Ratio",
                "download": "Download",
                "back": "Go Back"
            },
            "zh": {
                "title": "分析结果",
                "total_chars": "总汉字数",
                "highlighted_chars": "高亮字符数",
                "ratio": "比例",
                "download": "下载",
                "back": "返回"
            },
            "ja": {
                "title": "分析結果",
                "total_chars": "総漢字数",
                "highlighted_chars": "ハイライト文字数",
                "ratio": "比率",
                "download": "ダウンロード",
                "back": "戻る"
            }
        }
        
        texts = lang_texts.get(lang, lang_texts["ko"])
        
        # HTML 생성
        result_html = (
            f"<html><head><meta charset='utf-8'>"
            f"<style>"
            f"body {{ margin: 20px; font-family: '맑은 고딕', 'Malgun Gothic', sans-serif; font-size: 14px; line-height: 1.6; }}"
            f".highlight {{ color: red; font-weight: bold; }}"
            f".stats {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}"
            f".actions {{ margin: 20px 0; text-align: left; }}"
            f".btn {{ display: inline-block; padding: 10px 20px; margin: 5px; background: #007bff; color: white; text-decoration: none; border-radius: 5px; border: none; cursor: pointer; }}"
            f".btn:hover {{ background: #0056b3; }}"
            f".btn-secondary {{ background: #6c757d; }}"
            f".btn-secondary:hover {{ background: #545b62; }}"
            f".btn-primary {{ background: #007bff; }}"
            f".btn-primary:hover {{ background: #0056b3; }}"
            f".modal {{ display: none; position: fixed; z-index: 1000; left: 0; top: 0; width: 100%; height: 100%; background-color: rgba(0,0,0,0.5); }}"
            f".modal-dialog {{ margin: 5% auto; width: 400px; }}"
            f".modal-content {{ background-color: #fefefe; border: 1px solid #dee2e6; border-radius: 0.375rem; box-shadow: 0 0.5rem 1rem rgba(0, 0, 0, 0.15); }}"
            f".modal-header {{ display: flex; justify-content: space-between; align-items: center; padding: 1rem; border-bottom: 1px solid #dee2e6; border-top-left-radius: 0.375rem; border-top-right-radius: 0.375rem; }}"
            f".modal-title {{ margin: 0; font-size: 1.25rem; font-weight: 500; }}"
            f".btn-close {{ background: transparent; border: 0; font-size: 1.5rem; cursor: pointer; padding: 0; margin: 0; color: #aaa; }}"
            f".btn-close:hover {{ color: #000; }}"
            f".modal-body {{ padding: 1rem; }}"
            f".form-check {{ margin: 10px 0; }}"
            f".form-check-input {{ margin-right: 10px; }}"
            f".form-check-label {{ cursor: pointer; }}"
            f".modal-footer {{ padding: 1rem; border-top: 1px solid #dee2e6; text-align: right; }}"
            f".content {{ white-space: pre-wrap; word-wrap: break-word; }}"
            f"</style></head><body>"
        )
        
        # 통계 정보
        total_count = len(index_mapping)
        red_count = len(red_positions)
        ratio = (red_count / total_count * 100) if total_count > 0 else 0
        
        result_html += f"""
        <div class="stats">
            <h3>{texts['title']}</h3>
            <p><strong>{texts['total_chars']}:</strong> {total_count}자</p>
            <p><strong>{texts['highlighted_chars']}:</strong> {red_count}자</p>
            <p><strong>{texts['ratio']}:</strong> {ratio:.2f}%</p>
        </div>
        <div class="actions">
            <button onclick="showDownloadModal()" class="btn">{texts['download']}</button>
        </div>
        <div class="content">
        """
        
        # 텍스트 하이라이트
        for i, ch in enumerate(orig_text):
            if i in red_positions:
                result_html += f'<span class="highlight">{html_module.escape(ch)}</span>'
            else:
                # 줄바꿈과 탭을 HTML 엔티티로 변환
                if ch == '\n':
                    result_html += '<br>'
                elif ch == '\t':
                    result_html += '&nbsp;&nbsp;&nbsp;&nbsp;'  # 4개 공백
                else:
                    result_html += html_module.escape(ch)
        
        # 다운로드 모달과 스크립트 추가
        modal_html = """
        </div>

        <!-- Download Modal -->
        <div id="downloadModal" class="modal">
            <div class="modal-dialog">
                <div class="modal-content">
                    <div class="modal-header">
                        <h5 class="modal-title">다운로드 옵션</h5>
                        <button type="button" class="btn-close" onclick="closeDownloadModal()">&times;</button>
                    </div>
                    <div class="modal-body">
                        <div class="form-check">
                            <input class="form-check-input" type="radio" name="downloadFormat" id="formatHtml" value="html" checked>
                            <label class="form-check-label" for="formatHtml">HTML 파일 (.html)</label>
                        </div>
                        <div class="form-check">
                            <input class="form-check-input" type="radio" name="downloadFormat" id="formatDocx" value="docx">
                            <label class="form-check-label" for="formatDocx">Word 문서 (.docx)</label>
                        </div>
                        <div class="form-check">
                            <input class="form-check-input" type="radio" name="downloadFormat" id="formatHwp" value="hwp">
                            <label class="form-check-label" for="formatHwp">한글 문서 (.hwp)</label>
                        </div>
                    </div>
                    <div class="modal-footer">
                        <button type="button" class="btn btn-secondary" onclick="closeDownloadModal()">취소</button>
                        <button type="button" class="btn btn-primary" onclick="downloadFile()">다운로드</button>
                    </div>
                </div>
            </div>
        </div>

        <script>
        const highlightData = {
            file_name: '""" + file_name + """',
            selected_ngrams: '""" + ','.join(selected_ngrams) + """',
            lang: '""" + lang + """'
        };

        function showDownloadModal() {
            document.getElementById('downloadModal').style.display = 'block';
        }

        function closeDownloadModal() {
            document.getElementById('downloadModal').style.display = 'none';
        }

        async function downloadFile() {
            const format = document.querySelector('input[name="downloadFormat"]:checked').value;
            const { file_name, selected_ngrams, lang } = highlightData;
            
            try {
                const url = `/api/download-highlight?file_name=${encodeURIComponent(file_name)}&selected_ngrams=${encodeURIComponent(selected_ngrams)}&lang=${encodeURIComponent(lang)}&format_type=${encodeURIComponent(format)}`;
                
                const response = await fetch(url, {
                    method: 'GET'
                });

                if (response.ok) {
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `ngsm.${format}`;
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    window.URL.revokeObjectURL(url);
                } else {
                    const result = await response.json();
                    alert(result.detail || '다운로드 중 오류가 발생했습니다.');
                }
            } catch (error) {
                alert('다운로드 중 오류가 발생했습니다: ' + error.message);
            } finally {
                closeDownloadModal();
            }
        }

        // 모달 외부 클릭 시 닫기
        window.onclick = function(event) {
            const modal = document.getElementById('downloadModal');
            if (event.target == modal) {
                closeDownloadModal();
            }
        }
        </script>
        </body></html>"""
        
        result_html += modal_html
        return result_html

    def apply_to_original_text_as_text(self, file_name: str, selected_ngrams: Set[str], lang: str = "ko"):
        """원본 텍스트에 선택된 N-gram 하이라이트 적용 (텍스트 형식)"""
        if file_name not in self.original_text_data:
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        orig_text = self.original_text_data[file_name]
        processed_text = self.files_data[file_name]
        index_mapping = self.index_map[file_name]
        
        # 하이라이트할 위치 찾기
        red_positions = set()
        for ngram in selected_ngrams:
            start_idx = 0
            while True:
                pos = processed_text.find(ngram, start_idx)
                if pos == -1:
                    break
                for offset in range(len(ngram)):
                    if pos + offset < len(index_mapping):
                        red_positions.add(index_mapping[pos + offset])
                start_idx = pos + 1
        
        # 언어별 텍스트
        lang_texts = {
            "ko": {
                "title": "분석 결과",
                "total_chars": "전체 한자수",
                "highlighted_chars": "해당 글자수",
                "ratio": "비율"
            },
            "en": {
                "title": "Analysis Results",
                "total_chars": "Total Hanja Count",
                "highlighted_chars": "Highlighted Characters",
                "ratio": "Ratio"
            },
            "zh": {
                "title": "分析结果",
                "total_chars": "总汉字数",
                "highlighted_chars": "高亮字符数",
                "ratio": "比例"
            },
            "ja": {
                "title": "分析結果",
                "total_chars": "総漢字数",
                "highlighted_chars": "ハイライト文字数",
                "ratio": "比率"
            }
        }
        
        texts = lang_texts.get(lang, lang_texts["ko"])
        
        # 통계 정보
        total_count = len(index_mapping)
        red_count = len(red_positions)
        ratio = (red_count / total_count * 100) if total_count > 0 else 0
        
        # 텍스트 결과 생성
        result_text = f"{texts['title']}\n"
        result_text += f"{'='*50}\n"
        result_text += f"{texts['total_chars']}: {total_count}자\n"
        result_text += f"{texts['highlighted_chars']}: {red_count}자\n"
        result_text += f"{texts['ratio']}: {ratio:.2f}%\n"
        result_text += f"{'='*50}\n\n"
        
        # 텍스트 하이라이트 (대괄호로 표시)
        for i, ch in enumerate(orig_text):
            if i in red_positions:
                result_text += f"[{ch}]"
            else:
                # 원본 텍스트의 줄바꿈과 공백을 그대로 유지
                result_text += ch
        
        return result_text

    def apply_to_original_text_as_docx(self, file_name: str, selected_ngrams: Set[str], lang: str = "ko"):
        """원본 텍스트에 선택된 N-gram 하이라이트 적용 (Word 문서 형식)"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            from docx.oxml.ns import qn
        except ImportError:
            # python-docx가 설치되지 않은 경우 텍스트로 대체
            text_result = self.apply_to_original_text_as_text(file_name, selected_ngrams, lang)
            return text_result.encode('utf-8-sig')
        
        if file_name not in self.original_text_data:
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        orig_text = self.original_text_data[file_name]
        processed_text = self.files_data[file_name]
        index_mapping = self.index_map[file_name]
        
        # 하이라이트할 위치 찾기
        red_positions = set()
        for ngram in selected_ngrams:
            start_idx = 0
            while True:
                pos = processed_text.find(ngram, start_idx)
                if pos == -1:
                    break
                for offset in range(len(ngram)):
                    if pos + offset < len(index_mapping):
                        red_positions.add(index_mapping[pos + offset])
                start_idx = pos + 1
        
        # 언어별 텍스트
        lang_texts = {
            "ko": {
                "title": "분석 결과",
                "total_chars": "전체 한자수",
                "highlighted_chars": "해당 글자수",
                "ratio": "비율"
            },
            "en": {
                "title": "Analysis Results",
                "total_chars": "Total Hanja Count",
                "highlighted_chars": "Highlighted Characters",
                "ratio": "Ratio"
            },
            "zh": {
                "title": "分析结果",
                "total_chars": "总汉字数",
                "highlighted_chars": "高亮字符数",
                "ratio": "比例"
            },
            "ja": {
                "title": "分析結果",
                "total_chars": "総漢字数",
                "highlighted_chars": "ハイライト文字数",
                "ratio": "比率"
            }
        }
        
        texts = lang_texts.get(lang, lang_texts["ko"])
        
        # 통계 정보
        total_count = len(index_mapping)
        red_count = len(red_positions)
        ratio = (red_count / total_count * 100) if total_count > 0 else 0
        
        # Word 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading(texts['title'], 0)
        
        # 통계 정보 추가
        doc.add_paragraph(f"{texts['total_chars']}: {total_count}자")
        doc.add_paragraph(f"{texts['highlighted_chars']}: {red_count}자")
        doc.add_paragraph(f"{texts['ratio']}: {ratio:.2f}%")
        
        # 구분선 추가
        doc.add_paragraph("=" * 50)
        
        # 텍스트 내용 추가 (하이라이트 포함)
        content_para = doc.add_paragraph()
        
        for i, ch in enumerate(orig_text):
            if ch == '\n':
                # 줄바꿈 시 새 단락 생성
                content_para = doc.add_paragraph()
            elif ch == '\t':
                # 탭은 공백으로 대체
                run = content_para.add_run('    ')
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run = content_para.add_run(ch)
                if i in red_positions:
                    # 빨간색으로 하이라이트
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True
                else:
                    # 검은색으로 표시
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 한글 폰트 설정
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 바이트로 변환
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer.getvalue()

    def apply_to_original_text_as_hwp(self, file_name: str, selected_ngrams: Set[str], lang: str = "ko"):
        """원본 텍스트에 선택된 N-gram 하이라이트 적용 (한글 문서 형식)"""
        if file_name not in self.original_text_data:
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        orig_text = self.original_text_data[file_name]
        processed_text = self.files_data[file_name]
        index_mapping = self.index_map[file_name]
        
        # 하이라이트할 위치 찾기
        red_positions = set()
        for ngram in selected_ngrams:
            start_idx = 0
            while True:
                pos = processed_text.find(ngram, start_idx)
                if pos == -1:
                    break
                for offset in range(len(ngram)):
                    if pos + offset < len(index_mapping):
                        red_positions.add(index_mapping[pos + offset])
                start_idx = pos + 1
        
        # 언어별 텍스트
        lang_texts = {
            "ko": {
                "title": "분석 결과",
                "total_chars": "전체 한자수",
                "highlighted_chars": "해당 글자수",
                "ratio": "비율"
            },
            "en": {
                "title": "Analysis Results",
                "total_chars": "Total Hanja Count",
                "highlighted_chars": "Highlighted Characters",
                "ratio": "Ratio"
            },
            "zh": {
                "title": "分析结果",
                "total_chars": "总汉字数",
                "highlighted_chars": "高亮字符数",
                "ratio": "比例"
            },
            "ja": {
                "title": "分析結果",
                "total_chars": "総漢字数",
                "highlighted_chars": "ハイライト文字数",
                "ratio": "比率"
            }
        }
        
        texts = lang_texts.get(lang, lang_texts["ko"])
        
        # 통계 정보
        total_count = len(index_mapping)
        red_count = len(red_positions)
        ratio = (red_count / total_count * 100) if total_count > 0 else 0
        
        # HTML 형식으로 한글 문서 결과 생성 (빨간색 하이라이트 포함)
        result_html = f"""<html><head><meta charset='utf-8'>
        <style>
        body {{ font-family: '맑은 고딕', 'Malgun Gothic', sans-serif; margin: 20px; font-size: 14px; line-height: 1.6; }}
        .highlight {{ color: red; font-weight: bold; }}
        .stats {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
        .content {{ white-space: pre-wrap; word-wrap: break-word; }}
        </style></head><body>
        <div class="stats">
            <h3>{texts['title']}</h3>
            <p><strong>{texts['total_chars']}:</strong> {total_count}자</p>
            <p><strong>{texts['highlighted_chars']}:</strong> {red_count}자</p>
            <p><strong>{texts['ratio']}:</strong> {ratio:.2f}%</p>
        </div>
        <div class="content">
        """
        
        # 텍스트 하이라이트 (HTML 색상 태그 사용)
        for i, ch in enumerate(orig_text):
            if i in red_positions:
                result_html += f'<span class="highlight">{html_module.escape(ch)}</span>'
            else:
                # 줄바꿈과 탭을 HTML 엔티티로 변환
                if ch == '\n':
                    result_html += '<br>'
                elif ch == '\t':
                    result_html += '&nbsp;&nbsp;&nbsp;&nbsp;'  # 4개 공백
                else:
                    result_html += html_module.escape(ch)
        
        result_html += "</div></body></html>"
        
        return result_html.encode('utf-8-sig')

    def create_docx_from_results(self, results: List[str]) -> bytes:
        """분석 결과를 Word 문서로 변환"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            from docx.oxml.ns import qn
        except ImportError:
            # python-docx가 설치되지 않은 경우 텍스트로 대체
            text_result = "\n".join(results)
            return text_result.encode('utf-8-sig')
        
        # Word 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading("N-gram 분석 결과", 0)
        
        # 생성일 추가
        doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 구분선 추가
        doc.add_paragraph("=" * 50)
        
        # 분석 결과 추가
        for result in results:
            doc.add_paragraph(result)
        
        # 한글 폰트 설정
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 바이트로 변환
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer.getvalue()

# 전역 분석기 인스턴스
analyzer = NgramAnalyzer()

@app.get("/", response_class=HTMLResponse)
def home(request: Request, lang: str = "ko"):
    """메인 페이지"""
    if lang not in LOCALES:
        lang = "ko"
    return templates.TemplateResponse("index.html", {
        "request": request,
        "locale": LOCALES[lang],
        "current_lang": lang
    })

@app.post("/api/analyze")
async def analyze_files(
    min_n: int = Form(...),
    max_n: int = Form(...),
    files: List[UploadFile] = File(...)
):
    """파일 분석 API"""
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="최소 2개 이상의 파일을 선택해야 합니다.")
    
    try:
        analyzer.process_files(files, min_n, max_n)
        results = analyzer.get_analysis_results()
        
        return JSONResponse({
            "success": True,
            "filenames": list(analyzer.files_data.keys()),
            "results": results,
            "data_count": len(results)
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"분석 중 오류 발생: {str(e)}")

@app.post("/api/filter")
async def filter_ngrams(
    mode: str = Form(...),
    selected_files: str = Form(...),  # JSON string
    include_all_common: bool = Form(False),
    sort_option: str = Form("빈도수 오름차순")
):
    """N-gram 필터링 API"""
    try:
        if mode == "common":
            selected_files_list = json.loads(selected_files)
            results = analyzer.get_common_ngrams(selected_files_list, include_all_common)
            # 정렬 적용
            file_keys = selected_files_list
            sorted_ngrams = analyzer.sort_ngrams([line.split(":")[0].strip() for line in results], file_keys, sort_option)
            # 정렬된 결과 재구성
            sorted_results = []
            for ngram in sorted_ngrams:
                for line in results:
                    if line.startswith(ngram + ":"):
                        sorted_results.append(line)
                        break
            results = sorted_results
        else:
            results = analyzer.get_analysis_results(sort_option)
        
        return JSONResponse({
            "success": True,
            "results": results,
            "data_count": len(results)
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"필터링 중 오류 발생: {str(e)}")

@app.post("/api/wordcloud")
async def generate_wordcloud(
    ngrams_data: str = Form(...)  # JSON string of ngram results
):
    """워드클라우드 생성 API (N-gram.py 방식)"""
    try:
        # N-gram 데이터에서 빈도수 추출
        freq_dict = {}
        ngrams_list = json.loads(ngrams_data)
        for line in ngrams_list:
            if line.strip():
                parts = line.split(":")
                if len(parts) < 2:
                    continue
                ngram = parts[0].strip()
                counts = re.findall(r'\d+', line)
                total = sum(int(x) for x in counts) if counts else 0
                freq_dict[ngram] = total
        if not freq_dict:
            raise HTTPException(status_code=400, detail="워드클라우드 생성을 위한 유효한 데이터가 없습니다.")

        # 폰트 경로 자동 감지 (윈도우/리눅스)
        font_path = "static/fonts/NotoSansCJKtc-Regular.otf"
        if not os.path.exists(font_path):
            font_path = "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"
            if not os.path.exists(font_path):
                font_path = None

        # diverse_color_func (N-gram.py와 동일)
        def diverse_color_func(word, font_size, position, orientation, random_state=None, **kwargs):
            if random_state is None:
                random_state = random.Random()
            while True:
                hue = random_state.uniform(0, 360)
                if not (50 <= hue <= 70):
                    break
            h = hue / 360.0
            s = random_state.uniform(0.7, 1.0)
            l = random_state.uniform(0.3, 0.7)
            r, g, b = colorsys.hls_to_rgb(h, l, s)
            r, g, b = int(r*255), int(g*255), int(b*255)
            return f"#{r:02x}{g:02x}{b:02x}"

        # 워드클라우드 생성 (N-gram.py와 동일)
        wc = WordCloud(
            font_path=font_path,
            width=800,
            height=600,
            background_color="white",
            color_func=diverse_color_func
        )
        wc.generate_from_frequencies(freq_dict)

        # 이미지를 바이트로 변환
        img_buffer = io.BytesIO()
        wc.to_image().save(img_buffer, format="PNG")
        img_buffer.seek(0)

        return StreamingResponse(
            io.BytesIO(img_buffer.getvalue()),
            media_type="image/png",
            headers={"Content-Disposition": "attachment; filename=wordcloud.png"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"워드클라우드 생성 중 오류 발생: {str(e)}")

@app.post("/api/apply-highlight")
async def apply_highlight(
    file_name: str = Form(...),
    selected_ngrams: str = Form(...),  # JSON string
    lang: str = Form("ko")
):
    """원본 텍스트에 하이라이트 적용 API"""
    try:
        # JSON 파싱 오류 처리
        try:
            ngrams_set = set(json.loads(selected_ngrams))
        except json.JSONDecodeError as e:
            print(f"JSON 파싱 오류: {e}")
            raise HTTPException(status_code=400, detail="선택된 N-gram 데이터 형식이 올바르지 않습니다.")
        
        if not ngrams_set:
            raise HTTPException(status_code=400, detail="선택된 N-gram이 없습니다.")
        
        print(f"하이라이트 적용 시작: 파일={file_name}, N-gram 수={len(ngrams_set)}, 언어={lang}")
        
        html_result = analyzer.apply_to_original_text(file_name, ngrams_set, lang)
        
        print(f"하이라이트 적용 완료: HTML 길이={len(html_result)}")
        
        return HTMLResponse(html_result)
    except HTTPException:
        raise
    except Exception as e:
        print(f"하이라이트 적용 중 오류: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"하이라이트 적용 중 오류 발생: {str(e)}")

@app.get("/api/download-highlight")
async def download_highlight(
    file_name: str,
    selected_ngrams: str,
    lang: str = "ko",
    format_type: str = "html"
):
    """하이라이트된 결과 다운로드 API"""
    try:
        ngrams_set = set(selected_ngrams.split(','))
        
        if format_type == "html":
            html_result = analyzer.apply_to_original_text(file_name, ngrams_set, lang)
            
            return StreamingResponse(
                io.BytesIO(html_result.encode('utf-8-sig')),  # BOM 추가로 인코딩 문제 해결
                media_type="text/html; charset=utf-8",
                headers={"Content-Disposition": "attachment; filename=ngsm.html"}
            )
        elif format_type == "txt":
            # 텍스트 형식으로 다운로드
            text_result = analyzer.apply_to_original_text_as_text(file_name, ngrams_set, lang)
            
            return StreamingResponse(
                io.BytesIO(text_result.encode('utf-8-sig')),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": "attachment; filename=ngsm.txt"}
            )
        elif format_type == "docx":
            # Word 문서 형식으로 다운로드
            docx_result = analyzer.apply_to_original_text_as_docx(file_name, ngrams_set, lang)
            
            return StreamingResponse(
                io.BytesIO(docx_result),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=ngsm.docx"}
            )
        elif format_type == "hwp":
            # 한글 문서 형식으로 다운로드 (하이라이트 포함)
            hwp_result = analyzer.apply_to_original_text_as_hwp(file_name, ngrams_set, lang)
            
            return StreamingResponse(
                io.BytesIO(hwp_result),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": "attachment; filename=ngsm.hwp"}
            )
        else:
            raise HTTPException(status_code=400, detail="지원하지 않는 형식입니다.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"다운로드 중 오류 발생: {str(e)}")

@app.post("/api/download")
async def download_results(
    results_data: str = Form(...),  # JSON string
    format_type: str = Form("html")
):
    """분석 결과 다운로드 API"""
    try:
        results = json.loads(results_data)
        
        if format_type == "html":
            html_content = f"""
            <html><head><meta charset='utf-8'>
            <style>
            body {{ font-family: '맑은 고딕', 'Malgun Gothic', sans-serif; margin: 20px; }}
            .result {{ margin: 10px 0; padding: 5px; border-bottom: 1px solid #eee; }}
            .header {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
            </style></head><body>
            <div class="header">
                <h1>N-gram 분석 결과</h1>
                <p>생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
            """
            for result in results:
                html_content += f'<div class="result">{html_module.escape(result)}</div>'
            html_content += "</body></html>"
            
            return StreamingResponse(
                io.BytesIO(html_content.encode('utf-8-sig')),
                media_type="text/html; charset=utf-8",
                headers={"Content-Disposition": "attachment; filename*=UTF-8''ngram_analysis.html"}
            )
        elif format_type == "docx":
            # Word 문서 형식으로 다운로드
            docx_result = analyzer.create_docx_from_results(results)
            
            return StreamingResponse(
                io.BytesIO(docx_result),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename*=UTF-8''ngram_analysis.docx"}
            )
        elif format_type == "hwp":
            # 한글 문서 형식으로 다운로드 (텍스트로 대체)
            hwp_content = "N-gram 분석 결과\n"
            hwp_content += "=" * 50 + "\n"
            hwp_content += f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            hwp_content += "=" * 50 + "\n\n"
            for result in results:
                hwp_content += result + "\n"
            
            return StreamingResponse(
                io.BytesIO(hwp_content.encode('utf-8-sig')),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": "attachment; filename*=UTF-8''ngram_analysis.hwp"}
            )
        else:
            raise HTTPException(status_code=400, detail="지원하지 않는 형식입니다.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"다운로드 중 오류 발생: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("test_web:app", host="0.0.0.0", port=8000, reload=True) 
